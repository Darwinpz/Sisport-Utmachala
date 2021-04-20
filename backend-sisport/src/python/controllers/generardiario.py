from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#metodo de colores
def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

def crear_diario(ruta_destino,num_diario,periodo,tiempo,fecha,docente,tema,lista_contenidos,mi_objetivo,lista_actividades,lista_estrategias,mi_resumen,reflexion1,reflexion2,reflexion3,reflexion4):

	document = Document()

	#estilos generales del documento
	style = document.styles['Normal']
	font = style.font
	font.name = 'Arial'
	#font.size = Pt(12)
	#font.bold = True

	#colocar la imagen como encabezado
	header = document.sections[0].header
	p = header.add_paragraph()
	r = p.add_run()
	#r.add_picture('../principal/static/imagenes/encabezado.png')

	#creando tabla y darle color
	tabla=document.add_table(rows=5, cols=4, style='Table Grid')
	tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
	shade_cells([tabla.cell(0, 0), tabla.cell(0, 2), tabla.cell(1, 0), tabla.cell(2, 0), tabla.cell(3, 0), tabla.cell(4, 0)], "#ededed")

	#las primeras celdas de la tabla formato
	cell = tabla.cell(0, 0)
	cell.width = Inches(1.8)
	cell.text = 'CLASE N°:'+'\n'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(16)
	cell_font.bold=True

	cell = tabla.cell(0, 2)
	cell.text = 'PERIODO:'+'\n'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(16)
	cell_font.bold=True

	#contenido de las primeras celdas
	cell = tabla.cell(0,1)
	cell.width = Inches(0.5)
	cell.text = num_diario
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(14)
	cell_font.bold=True
	cell_font.color.rgb = RGBColor(196, 89, 17)

	cell = tabla.cell(0,3)
	cell.text = periodo
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(14)
	cell_font.bold=True
	cell_font.color.rgb = RGBColor(196, 89, 17)

	#el resto de celdas formato
	cell = tabla.cell(1,0)
	cell.text = '\n'+'TIEMPO:'+'\n'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(12)
	cell_font.bold=True

	cell = tabla.cell(2,0)
	cell.text = '\n'+'FECHA:'+'\n'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(12)
	cell_font.bold=True

	cell = tabla.cell(3,0)
	cell.text = '\n'+'DOCENTE GUÍA:'+'\n'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(12)
	cell_font.bold=True

	cell = tabla.cell(4,0)
	cell.text = '\n'+'TEMA DISCUTIDO:'+'\n'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size = Pt(12)
	cell_font.bold=True


	#celdas unidas del contenido
	heading_cells = tabla.rows[1].cells
	nueva_celda=heading_cells[1].merge(heading_cells[2]).merge(heading_cells[3])
	nueva_celda.text='\n'+tiempo+'\n'
	nueva_celda_font = nueva_celda.paragraphs[0].runs[0].font
	nueva_celda_font.size = Pt(12)


	heading_cells = tabla.rows[2].cells
	nueva_celda=heading_cells[1].merge(heading_cells[2]).merge(heading_cells[3])
	nueva_celda.text='\n'+fecha+'\n'
	nueva_celda_font = nueva_celda.paragraphs[0].runs[0].font
	nueva_celda_font.size = Pt(12)


	heading_cells = tabla.rows[3].cells
	nueva_celda=heading_cells[1].merge(heading_cells[2]).merge(heading_cells[3])
	nueva_celda.text='\n'+docente+'\n'
	nueva_celda_font = nueva_celda.paragraphs[0].runs[0].font
	nueva_celda_font.size = Pt(12)

	heading_cells = tabla.rows[4].cells
	nueva_celda=heading_cells[1].merge(heading_cells[2]).merge(heading_cells[3])
	nueva_celda.text='\n'+tema+'\n'
	nueva_celda_font = nueva_celda.paragraphs[0].runs[0].font
	nueva_celda_font.size = Pt(12)


	#celdas de abajo en el centro
	tabla.cell(1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabla.cell(2,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabla.cell(3,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabla.cell(4,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

	#contenido
	contenido=document.add_paragraph('\n')
	contenido.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	font.size = Pt(12)
	estilo_contenido=contenido.add_run('\n'+'Contenidos:')
	estilo_contenido.bold=True

	#lista_contenido=["hola","soy","un","contenido"]

	for i in lista_contenidos:

		p = document.add_paragraph()
		p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
		p.style = 'List Bullet'

		r = p.add_run()
		r.add_text(i)

	#objetivo
	objetivo=document.add_paragraph()
	objetivo.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	estilo_objetivo=objetivo.add_run('\n'+'Objetivo:')
	estilo_objetivo.bold=True

	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	p.style = 'List Bullet'

	r = p.add_run()
	r.add_text(mi_objetivo)

	#datos interesantes discutidos
	datos=document.add_paragraph()
	estilo_datos=datos.add_run('\n'+'Datos interesantes discutidos:'+'\n')
	estilo_datos.bold=True
	estilo_datos=datos.add_run('\n'+'Actividades durante la clase:')
	estilo_datos.bold=True

	#lista_actividades=["hola","soy","una","actividad"]
	for i in lista_actividades:

		p = document.add_paragraph()
		p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
		p.style = 'List Bullet'

		r = p.add_run()
		r.add_text(i)

	#estrategias
	estrategias=document.add_paragraph()
	estrategias.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	estilo_estrategias=estrategias.add_run('\n'+'Estrategias de aprendizaje:')
	estilo_estrategias.bold=True

	#lista_estrategias=["hola","soy","una","estrategia"]
	for i in lista_estrategias:

		p = document.add_paragraph()
		p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
		p.style = 'List Bullet'

		r = p.add_run()
		r.add_text(i)

	#resumen conceptual
	resumen=document.add_paragraph()
	resumen.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	estilo_resumen=resumen.add_run('\n'+'Resumen conceptual:')
	estilo_resumen.bold=True
	concepto=document.add_paragraph()
	concepto.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	estilo_concepto=concepto.add_run('\n'+mi_resumen+'\n')

	#reflexion
	reflexion=document.add_paragraph()
	reflexion.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	estilo_reflexion=reflexion.add_run('\n'+'Reflexionar:')
	estilo_reflexion.bold=True
	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	p.style = 'List Bullet'
	r = p.add_run()
	r.add_text('¿Qué cosas fueron difíciles?'+'\n')
	r.bold=True
	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	texto=p.add_run(reflexion1)

	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	p.style = 'List Bullet'
	r = p.add_run()
	r.add_text('¿Cuáles fueron fáciles?'+'\n')
	r.bold=True
	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	texto=p.add_run(reflexion2)

	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	p.style = 'List Bullet'
	r = p.add_run()
	r.add_text('¿Por qué?'+'\n')
	r.bold=True
	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	texto=p.add_run(reflexion3)

	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	p.style = 'List Bullet'
	r = p.add_run()
	r.add_text('¿Qué aprendí hoy?'+'\n')
	r.bold=True
	p = document.add_paragraph()
	p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	texto=p.add_run(reflexion4)

	#creando documento
	document.save(ruta_destino)