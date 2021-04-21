from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#metodo de colores
def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)

def crear_informe(ruta_destino,estudiante,asignatura,curso,profesor,periodo,informe):

	document = Document()

	#estilos generales del documento
	style = document.styles['Normal']
	font = style.font
	font.name = 'Arial'
	#font.size = Pt(12)
	#font.bold = True

	#colocar la imagen como encabezado
	header = document.sections[0].header
	p = header.add_paragraph()
	r = p.add_run()
	#r.add_picture('../principal/static/imagenes/encabezado.png')


	#descripcion
	descripcion=document.add_paragraph()
	descripcion.alignment=WD_ALIGN_PARAGRAPH.CENTER
	font.size = Pt(12)
	estilo_descripcion=descripcion.add_run('Informe final del cumplimiento de las expectativas al finalizar el curso')
	estilo_descripcion.bold=True

	#creando tabla y darle color
	tabla=document.add_table(rows=5, cols=2, style='Table Grid')
	tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
	shade_cells([tabla.cell(0, 0), tabla.cell(1, 0), tabla.cell(2, 0), tabla.cell(3, 0), tabla.cell(4, 0), tabla.cell(4, 0)],"#ededed")

	#tabla columna 1
	cell=tabla.cell(0,0)
	cell.text ='Estudiante:'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=True

	cell=tabla.cell(1,0)
	cell.text ='Asignatura:'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=True

	cell=tabla.cell(2,0)
	cell.text ='Curso:'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=True

	cell=tabla.cell(3,0)
	cell.text ='Profesor:'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=True

	cell=tabla.cell(4,0)
	cell.text ='Periodo Academico:'
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=True

	#tabla columna 2
	cell=tabla.cell(0,1)
	cell.text =estudiante
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=False

	cell=tabla.cell(1,1)
	cell.text =asignatura
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=False

	cell=tabla.cell(2,1)
	cell.text =curso
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=False

	cell=tabla.cell(3,1)
	cell.text =profesor
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=False

	cell=tabla.cell(4,1)
	cell.text =periodo
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.size=Pt(12)
	cell_font.bold=False

	#dar medida a la columna 1
	tabla.columns[0].width = Inches(3)

	#contenido
	contenido=document.add_paragraph('\n')
	contenido.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
	font.size = Pt(12)
	contenido_descripcion=contenido.add_run(informe)
	contenido_descripcion.bold=False
	

	#creando documento
	document.save(ruta_destino)
